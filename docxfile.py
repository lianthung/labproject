from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import os
import subprocess
from database import load_patient_from_db
from datetime import datetime

def open_file(pid):
    doc = Document()
    doc.add_heading('CROSS IMAGING', 0)
    doc.add_heading('DAWRPUI, NEAR SBI, AIZAWL', 2)
    doc.add_heading('ULTRASOUND REPORT', 2)

    patient = load_patient_from_db(pid)
    patient_sec = doc.add_paragraph(patient['pname']+' '+str(patient['age'])+'YR '+patient['address']+'\n')
    patient_sec.add_run('Ref Dr.:'+patient['ref_doctor']+'\nMobile:'+patient['mobile'])
    patient_sec.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    test_date = patient['test_date'].strftime('%Y%m%d')
    filename = test_date+'-2-'+patient['pname']+'.docx'
    doc.save(filename)

    os.startfile(filename)
    #subprocess.Popen('C:\Program Files (x86)\Microsoft Office\Office12\winword.exe')

